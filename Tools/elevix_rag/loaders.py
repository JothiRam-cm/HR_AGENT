import os
import pandas as pd
import docx
from typing import List, Dict, Any
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    BSHTMLLoader
)

class UnifiedLoader:
    def __init__(self):
        pass

    def load_pdf(self, file_path: str) -> List[Document]:
        """
        PDF: Page / section-based extraction.
        Keep page number in metadata.
        Metadata: {source, page_number}
        """
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        processed_docs = []
        for i, doc in enumerate(docs):
            # Normalize metadata
            doc.metadata = {
                "source_file": os.path.basename(file_path),
                "file_type": "pdf",
                "page_number": i + 1,
                "section": f"Page {i + 1}"
            }
            processed_docs.append(doc)
        return processed_docs

    def load_docx(self, file_path: str) -> List[Document]:
        """
        DOCX: Heading-aware parsing.
        Preserve headings. Each heading = one semantic unit.
        Content under heading becomes chunk.
        Metadata: {source, section_heading}
        """
        doc = docx.Document(file_path)
        documents = []
        current_section = "Initial Information"
        current_content = []

        for paragraph in doc.paragraphs:
            # Check if it's a heading
            is_heading = False
            if paragraph.style.name.startswith('Heading'):
                is_heading = True
            
            if is_heading:
                if current_content:
                    text = "\n".join(current_content).strip()
                    if text:
                        documents.append(Document(
                            page_content=text,
                            metadata={
                                "source_file": os.path.basename(file_path),
                                "file_type": "docx",
                                "section_heading": current_section,
                                "section": current_section
                            }
                        ))
                current_section = paragraph.text
                current_content = []
            else:
                current_content.append(paragraph.text)

        # Append last section
        if current_content:
            text = "\n".join(current_content).strip()
            if text:
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source_file": os.path.basename(file_path),
                        "file_type": "docx",
                        "section_heading": current_section,
                        "section": current_section
                    }
                ))
        return documents

    def load_txt(self, file_path: str) -> List[Document]:
        """
        TXT: Split by blank lines. Merge very small fragments.
        Metadata: {source, section}
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        paragraphs = content.split('\n\n')
        documents = []
        
        # Merge very small fragments (< 50 chars) with next one
        merged_paragraphs = []
        temp = ""
        for p in paragraphs:
            p = p.strip()
            if not p: continue
            if len(p) < 50:
                temp += p + "\n"
            else:
                merged_paragraphs.append(temp + p)
                temp = ""
        if temp:
            if merged_paragraphs:
                merged_paragraphs[-1] += "\n" + temp
            else:
                merged_paragraphs.append(temp)

        for i, p in enumerate(merged_paragraphs):
            documents.append(Document(
                page_content=p,
                metadata={
                    "source_file": os.path.basename(file_path),
                    "file_type": "txt",
                    "section": f"Paragraph {i + 1}"
                }
            ))
        return documents

    def load_md(self, file_path: str) -> List[Document]:
        """
        MD: Header-based sections.
        Each heading = one semantic unit.
        Content under heading becomes chunk.
        Metadata: {source, section_heading}
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
        
        documents = []
        current_header = "Intro"
        current_content = []
        
        for line in lines:
            if line.startswith('#'):
                if current_content:
                    text = "".join(current_content).strip()
                    if text:
                        documents.append(Document(
                            page_content=text,
                            metadata={
                                "source_file": os.path.basename(file_path),
                                "file_type": "md",
                                "section_heading": current_header,
                                "section": current_header
                            }
                        ))
                current_header = line.strip('# \n')
                current_content = []
            else:
                current_content.append(line)
        
        if current_content:
            text = "".join(current_content).strip()
            if text:
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source_file": os.path.basename(file_path),
                        "file_type": "md",
                        "section_heading": current_header,
                        "section": current_header
                    }
                ))
        return documents

    def load_html(self, file_path: str) -> List[Document]:
        """
        HTML: Tag-aware text extraction.
        Using BSHTMLLoader for tag awareness.
        Metadata: {source, section}
        """
        loader = BSHTMLLoader(file_path)
        docs = loader.load()
        for doc in docs:
            doc.metadata = {
                "source_file": os.path.basename(file_path),
                "file_type": "html",
                "section": "Body Content"
            }
        return docs

    def load_csv(self, file_path: str) -> List[Document]:
        """
        CSV: Row-based documents.
        """
        df = pd.read_csv(file_path)
        return self._process_dataframe(df, file_path, "csv")

        """
        Excel: Sheet -> row-based documents.
        Each sheet handled separately.
        """
        xls = None
        try:
            xls = pd.ExcelFile(file_path)
            all_docs = []
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                sheet_docs = self._process_dataframe(df, file_path, "excel", sheet_name=sheet_name)
                all_docs.extend(sheet_docs)
            return all_docs
        except Exception as e:
            print(f"Error loading Excel {file_path}: {e}")
            return []
        finally:
            if xls:
                xls.close()

    def _process_dataframe(self, df: pd.DataFrame, file_path: str, file_type: str, sheet_name: str = None) -> List[Document]:
        """
        Critical Logic:
        One row = one fact.
        ColumnName: Value format.
        Skip null values.
        Extract schema as separate document.
        """
        documents = []
        file_name = os.path.basename(file_path)
        
        # 1. Schema Awareness: Extract column names and inferred data types
        schema_info = []
        for col in df.columns:
            dtype = str(df[col].dtype)
            schema_info.append(f"- {col} ({dtype})")
        
        schema_content = f"Schema for {file_name}" + (f" (Sheet: {sheet_name})" if sheet_name else "") + ":\n"
        schema_content += "Columns:\n" + "\n".join(schema_info)
        
        schema_metadata = {
            "type": "schema",
            "source_file": file_name,
            "file_type": file_type,
            "section": "Schema"
        }
        if sheet_name:
            schema_metadata["sheet_name"] = sheet_name
            
        documents.append(Document(page_content=schema_content, metadata=schema_metadata))
        
        # 2. Row Conversion Logic
        for idx, row in df.iterrows():
            row_parts = []
            for col in df.columns:
                val = row[col]
                if pd.notna(val) and str(val).strip().lower() not in ["nan", "none", "null"]:
                    row_parts.append(f"{col}: {val}")
            
            row_content = "\n".join(row_parts)
            if row_content.strip():
                metadata = {
                    "source_file": file_name,
                    "file_type": file_type,
                    "row_index": idx,
                    "section": f"Row {idx}"
                }
                if sheet_name:
                    metadata["sheet_name"] = sheet_name
                
                documents.append(Document(page_content=row_content, metadata=metadata))
        
        return documents

    def load(self, file_path: str) -> List[Document]:
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".pdf":
                return self.load_pdf(file_path)
            elif ext == ".docx":
                return self.load_docx(file_path)
            elif ext in [".txt", ".text"]:
                return self.load_txt(file_path)
            elif ext == ".md":
                return self.load_md(file_path)
            elif ext in [".html", ".htm"]:
                return self.load_html(file_path)
            elif ext == ".csv":
                return self.load_csv(file_path)
            elif ext in [".xlsx", ".xls"]:
                return self.load_excel(file_path)
            else:
                print(f"Skipping unsupported file type: {ext} ({file_path})")
                return []
        except Exception as e:
            print(f"Error loading {file_path}: {e}")
            return []
