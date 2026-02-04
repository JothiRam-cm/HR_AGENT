import os
import pandas as pd
from docx import Document
from config import Config

def generate_all_mock_data():
    path = Config.DATA_PATH
    os.makedirs(path, exist_ok=True)
    
    # 1. Create CSV
    csv_data = {
        "employee_id": [101, 102, 103],
        "name": ["Alice Smith", "Bob Johnson", "Charlie Brown"],
        "department": ["HR", "Engineering", "Sales"],
        "salary": [75000, 95000, 80000],
        "bonus": [5000, 8000, 10000],
        "join_date": ["2022-01-15", "2021-06-20", "2023-03-10"]
    }
    df = pd.DataFrame(csv_data)
    df.to_csv(os.path.join(path, "employees.csv"), index=False)
    print(f"Mock CSV created at {os.path.join(path, 'employees.csv')}")

    # 2. Create MD
    md_content = """# Company Benefits Overview

## Health Insurance
We provide comprehensive health insurance.
- Dental covered 80%
- Vision covered 100%

## Retirement Plan
401k match up to 5% of salary.

## Vacation
3 weeks per year for new hires.
"""
    with open(os.path.join(path, "benefits.md"), "w") as f:
        f.write(md_content)
    print(f"Mock MD created at {os.path.join(path, 'benefits.md')}")

    # 3. Create DOCX
    doc = Document()
    doc.add_heading('Travel Policy', 0)
    doc.add_heading('1. Domestic Travel', level=1)
    doc.add_paragraph('Employees can book economy class for flights under 4 hours.')
    doc.add_heading('2. International Travel', level=1)
    doc.add_paragraph('Business class is allowed for flights over 8 hours.')
    doc.save(os.path.join(path, "travel_policy.docx"))
    print(f"Mock DOCX created at {os.path.join(path, 'travel_policy.docx')}")

    # 4. Create TXT
    txt_content = """General Office Hours
    
    The office is open from 8:00 AM to 6:00 PM.
    
    Security procedures:
    Always wear your badge.
    Guest must sign in at front desk.
    """
    with open(os.path.join(path, "office_info.txt"), "w") as f:
        f.write(txt_content)
    print(f"Mock TXT created at {os.path.join(path, 'office_info.txt')}")

if __name__ == "__main__":
    generate_all_mock_data()
